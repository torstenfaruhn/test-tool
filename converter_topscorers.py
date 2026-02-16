"""
Amateurvoetbal topscorers: tekstbestand (.txt/.docx) -> Word (.docx)

Doel (zoals TOPSCORERS voorbeeld.docx):
- Per sectie (kop met 'KLASSE' of 'DIVISIE') start de genummerde lijst opnieuw bij 1.
- Binnen een sectie: elk nieuw doelpunten-aantal => nieuw lijstnummer (Enter = nieuwe paragraaf).
- Spelers met hetzelfde aantal doelpunten => binnen hetzelfde nummer met Shift+Enter (soft line break).
- Alleen het lijstnummer (bijv. "1.") is vet; alle tekst in het item is regular.
- Volgorde uit de bron blijft exact behouden (geen sortering).
"""

from __future__ import annotations

import re
import tempfile
from io import BytesIO
from typing import List, Tuple, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Legacy: sommige bronnen beginnen items met "1." (optioneel)
NUMBER_RE = re.compile(r"^\s*\d+\.\s*")

# Herken doelpunten-aantal op regels die een nieuw item moeten starten
# Voorbeelden die matchen:
# "... - 14 doelpunten"
# "...-14 doelpunt"
GOALS_RE = re.compile(r"\s*-\s*(\d+)\s+doelpunt(?:en)?\s*$", re.IGNORECASE)


# ----------------------------
# Parsing helpers
# ----------------------------
def looks_like_player_stat_line(line: str) -> bool:
    s = line.strip()
    lower = s.lower()
    # Heuristiek: spelerregels hebben vaak "(...)"
    if "(" in s and ")" in s:
        return True
    # Of bevatten "- <getal> doelpunt(en)"
    if GOALS_RE.search(s):
        return True
    # Of bevatten "-" plus een getal en het woord "doelpunt"
    if "-" in s and re.search(r"\b\d+\b", s) and "doelpunt" in lower:
        return True
    return False


def is_section_heading(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    # Als de regel begint met "1." is het geen sectiekop
    if NUMBER_RE.match(s):
        return False
    upper = s.upper()
    # Sectiekoppen bevatten KLASSE of DIVISIE
    if "KLASSE" not in upper and "DIVISIE" not in upper:
        return False
    # Maar spelerregels kunnen ook 'divisie/klasse' bevatten (bv. in clubnaam/omschrijving)
    if looks_like_player_stat_line(s):
        return False
    return True


def strip_source_rank_number(line: str) -> str:
    return re.sub(r"^\s*\d+\.\s*", "", line, count=1)


def parse_sections(text: str) -> List[Tuple[str, List[List[str]]]]:
    """
    Parseer inputtekst naar secties (titel + groepen regels).

    Belangrijk:
    - Headings: regels met 'KLASSE' of 'DIVISIE'.
    - Een regel met '- <N> doelpunt(en)' start (of wisselt) de huidige goals-groep.
    - Regels zonder goals-suffix horen bij de laatst geziene goals-groep (zelfde lijstnummer).
    - Als de bron toch met '1.' werkt, ondersteunen we dat ook:
      een nieuwe '1.' start altijd een nieuwe groep (nieuw lijstnummer).

    De volgorde uit de bron blijft exact behouden.
    """
    lines = text.splitlines()

    sections: List[Tuple[str, List[List[str]]]] = []
    current_title: Optional[str] = None
    current_groups: List[List[str]] = []
    current_group: List[str] = []
    current_goals: Optional[int] = None

    def flush_group() -> None:
        nonlocal current_group, current_groups
        if current_group:
            current_groups.append(current_group)
            current_group = []

    def flush_section() -> None:
        nonlocal current_title, current_groups, sections, current_goals
        if current_title and current_groups:
            sections.append((current_title, current_groups))
        current_groups = []
        current_goals = None

    for raw in lines:
        line = raw.strip()
        if not line:
            continue

        if is_section_heading(line):
            flush_group()
            flush_section()
            current_title = line
            continue

        # 1) Legacy input: "1. ..." start altijd een nieuwe groep
        if NUMBER_RE.match(line):
            flush_group()
            stripped = strip_source_rank_number(line)
            current_group = [stripped]
            m = GOALS_RE.search(stripped)
            current_goals = int(m.group(1)) if m else None
            continue

        # 2) Huidige input: groepeer op goals-wissel
        m = GOALS_RE.search(line)
        if m:
            goals = int(m.group(1))
            if current_goals is None:
                # Eerste goals in deze sectie
                flush_group()
                current_group = [line]
                current_goals = goals
            elif goals != current_goals:
                # Nieuw aantal goals => nieuw lijstitem (Enter)
                flush_group()
                current_group = [line]
                current_goals = goals
            else:
                # Zelfde goals, maar deze regel bevat ook het suffix => dezelfde groep
                if not current_group:
                    current_group = [line]
                else:
                    current_group.append(line)
        else:
            # Geen goals suffix: hoort bij huidige groep (Shift+Enter)
            if not current_group:
                # Defensief: start toch een groep zodat we geen regels verliezen
                current_group = [line]
            else:
                current_group.append(line)

    flush_group()
    flush_section()
    return sections


# ----------------------------
# DOCX numbering (robuust: nieuw numId per sectie)
# ----------------------------
def _new_numbering_numid_for_section(doc: Document, bold_number: bool = True) -> int:
    """
    Maak een nieuwe numbering instance (numId) die start bij 1.
    Robuust: elk numId is een aparte lijst (herstart per sectie).

    bold_number=True zet alleen het nummer vet via numbering rPr.
    """
    numbering_part = doc.part.numbering_part
    numbering = numbering_part.numbering_definitions._numbering  # CT_Numbering

    # Bepaal nieuwe IDs
    existing_abs = [
        int(n.get(qn("w:abstractNumId")))
        for n in numbering.findall(qn("w:abstractNum"))
        if n.get(qn("w:abstractNumId")) is not None
    ]
    existing_num = [
        int(n.get(qn("w:numId")))
        for n in numbering.findall(qn("w:num"))
        if n.get(qn("w:numId")) is not None
    ]
    abstract_id = (max(existing_abs) + 1) if existing_abs else 1
    num_id = (max(existing_num) + 1) if existing_num else 1

    # <w:abstractNum w:abstractNumId="...">
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    # <w:multiLevelType w:val="singleLevel"/>
    mlt = OxmlElement("w:multiLevelType")
    mlt.set(qn("w:val"), "singleLevel")
    abstract.append(mlt)

    # <w:lvl w:ilvl="0">
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    # Start bij 1
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    # Decimal numbering
    numfmt = OxmlElement("w:numFmt")
    numfmt.set(qn("w:val"), "decimal")
    lvl.append(numfmt)

    # Tekst van het nummer: "%1."
    lvltext = OxmlElement("w:lvlText")
    lvltext.set(qn("w:val"), "%1.")
    lvl.append(lvltext)

    # Spatie na nummer
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    lvl.append(suff)

    # Alleen het nummer vet maken
    if bold_number:
        rpr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        b.set(qn("w:val"), "1")
        rpr.append(b)
        lvl.append(rpr)

    # Inspringing (net als standaard lijst)
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")     # ~0.5 inch
    ind.set(qn("w:hanging"), "360")  # hangend
    ppr.append(ind)
    lvl.append(ppr)

    abstract.append(lvl)
    numbering.append(abstract)

    # <w:num w:numId="..."><w:abstractNumId w:val="..."/></w:num>
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    absref = OxmlElement("w:abstractNumId")
    absref.set(qn("w:val"), str(abstract_id))
    num.append(absref)
    numbering.append(num)

    return num_id


def _apply_numid_to_paragraph(paragraph, num_id: int, ilvl: int = 0) -> None:
    """Koppel een paragraaf aan een nummering (numId) op level ilvl."""
    p = paragraph._p
    ppr = p.get_or_add_pPr()

    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)

    ilvl_el = numpr.find(qn("w:ilvl"))
    if ilvl_el is None:
        ilvl_el = OxmlElement("w:ilvl")
        numpr.append(ilvl_el)
    ilvl_el.set(qn("w:val"), str(ilvl))

    numid_el = numpr.find(qn("w:numId"))
    if numid_el is None:
        numid_el = OxmlElement("w:numId")
        numpr.append(numid_el)
    numid_el.set(qn("w:val"), str(num_id))


# ----------------------------
# DOCX output
# ----------------------------
def topscorers_text_to_docx_bytes(text: str) -> bytes:
    """
    Converteer topscorers tekst naar docx-bytes.

    Structuur:
    - Sectiekop: aparte paragraaf
    - Daarna: per goals-groep één genummerde paragraaf (Enter => nieuw item)
      - group[0] als eerste regel (regular tekst)
      - group[1:] als extra regels met Shift+Enter binnen hetzelfde item
    """
    doc = Document()
    sections = parse_sections(text)

    for title, groups in sections:
        # Sectiekop
        doc.add_paragraph(title, style="Heading 3")

        # Robuuste herstart: nieuw numId per sectie
        num_id = _new_numbering_numid_for_section(doc, bold_number=True)

        # Eén paragraaf per groep (nieuw nummer per nieuw goals-aantal)
        for group in groups:
            if not group:
                continue

            p = doc.add_paragraph()
            _apply_numid_to_paragraph(p, num_id, ilvl=0)

            # Eerste regel (regular)
            p.add_run(group[0])

            # Extra regels (zelfde nummer) => Shift+Enter
            for extra_line in group[1:]:
                br = p.add_run()
                br.add_break()  # soft line break (Shift+Enter)
                p.add_run(extra_line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ----------------------------
# Input lezen (txt/docx upload)
# ----------------------------
def extract_text_from_upload(raw: bytes, filename: str) -> str:
    """Lees upload (.txt of .docx) en geef de tekstinhoud terug."""
    name = (filename or "").lower()
    if name.endswith(".docx"):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
            tmp.write(raw)
            tmp.flush()
            doc = Document(tmp.name)

            lines: List[str] = []
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    lines.append(t)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for ln in cell.text.splitlines():
                            t = ln.strip()
                            if t:
                                lines.append(t)

            return "\n".join(lines)

    # default: txt
    return raw.decode("utf-8", errors="replace")


# ------------------------------------------------------------
# Compat: bytes-uploader API (gebruikt door app.py)
# ------------------------------------------------------------
def extract_text_from_upload_bytes(raw: bytes, filename: str) -> str:
    """Lees upload-bytes en geef tekst terug (stabiele API voor app.py)."""
    return extract_text_from_upload(raw, filename)


# ------------------------------------------------------------
# Backward compat (optioneel): oude HTML functie is verwijderd
# ------------------------------------------------------------
def topscorers_text_to_cueweb_html(text: str) -> str:
    raise ImportError(
        "topscorers_text_to_cueweb_html is verwijderd. Gebruik topscorers_text_to_docx_bytes."
    )
