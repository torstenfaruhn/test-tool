"""Amateurvoetbal online: Cue Print (txt met tags) -> Word (.docx)

Deze module volgt de logica uit:
`03 colab_tag_converter-docx.ipynb`

Afspraken:
- Input bevat blokken:
  - 1x <subhead_lead> (divisie/klasse)
  - gevolgd door meerdere items: <subhead> en optioneel <howto_facts>
- Output:
  - <subhead_lead> wordt 1x als kopregel geplaatst (BOLD + UPPERCASE)
  - Elke <subhead> is 1 alinea (wedstrijdregel)
  - Als <howto_facts> tekst heeft: op volgende regel in dezelfde alinea met Shift+Enter, altijd italic
  - Lege <howto_facts>: overslaan
  - Tussen competitieblokken: 1 lege alinea

Privacy/logging:
- Geen logging van artikelteksten of persoonsdata; alleen technische status/tellingen (caller bepaalt).
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_BREAK

_TAG_PATTERN = re.compile(
    r"<(subhead_lead|subhead|howto_facts)>(.*?)</\1>",
    re.DOTALL | re.IGNORECASE,
)


@dataclass(frozen=True)
class _Token:
    kind: str
    text: str


@dataclass(frozen=True)
class _Item:
    header: str
    subhead: str
    facts: Optional[str]  # None als leeg/ontbrekend


def _extract_tokens(raw: str) -> List[_Token]:
    tokens: List[_Token] = []
    for m in _TAG_PATTERN.finditer(raw or ""):
        kind = (m.group(1) or "").lower().strip()
        text = (m.group(2) or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        tokens.append(_Token(kind=kind, text=text))
    return tokens


def _tokens_to_items(tokens: List[_Token]) -> Tuple[List[_Item], Dict[str, int]]:
    items: List[_Item] = []
    current_header: Optional[str] = None

    stats = {
        "tokens_total": len(tokens),
        "headers_seen": 0,
        "items_total": 0,
        "items_with_facts": 0,
        "empty_facts_skipped": 0,
    }

    i = 0
    while i < len(tokens):
        t = tokens[i]

        if t.kind == "subhead_lead":
            current_header = t.text
            stats["headers_seen"] += 1
            i += 1
            continue

        if t.kind == "subhead":
            if not current_header:
                # Zonder header geen geldig blok; sla veilig over
                i += 1
                continue

            subhead_text = t.text
            facts_text: Optional[str] = None

            if i + 1 < len(tokens) and tokens[i + 1].kind == "howto_facts":
                candidate = tokens[i + 1].text
                if candidate.strip():
                    facts_text = candidate.strip()
                    stats["items_with_facts"] += 1
                else:
                    stats["empty_facts_skipped"] += 1
                i += 2
            else:
                # Geen facts-tag
                stats["empty_facts_skipped"] += 1
                i += 1

            items.append(_Item(header=current_header, subhead=subhead_text, facts=facts_text))
            stats["items_total"] += 1
            continue

        # Losse howto_facts of andere rommel: overslaan
        i += 1

    return items, stats


def cueprint_txt_to_docx_bytes(text_in: str) -> Tuple[bytes, Dict[str, int]]:
    """Converteer Cue Print-tagtekst naar .docx-bytes.

    Returns:
        (docx_bytes, stats) — stats bevat alleen technische tellingen.
    """
    tokens = _extract_tokens(text_in)
    items, stats = _tokens_to_items(tokens)

    doc = Document()
    prev_header: Optional[str] = None
    block_separators_added = 0
    headers_placed = 0

    for it in items:
        header = (it.header or "").strip()

        # Nieuw competitieblok?
        if header and header != prev_header:
            if prev_header is not None:
                doc.add_paragraph("")  # 1 lege alinea tussen blokken
                block_separators_added += 1

            hp = doc.add_paragraph()
            hr = hp.add_run(header.upper())
            hr.bold = True
            headers_placed += 1
            prev_header = header

        # Wedstrijd-item (1 alinea)
        p = doc.add_paragraph()
        r1 = p.add_run((it.subhead or "").strip())

        facts = (it.facts or "").strip()
        if facts:
            r1.add_break(WD_BREAK.LINE)  # Shift+Enter
            r2 = p.add_run(facts)
            r2.italic = True

    stats.update(
        {
            "headers_placed": headers_placed,
            "block_separators_added": block_separators_added,
        }
    )

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue(), stats
