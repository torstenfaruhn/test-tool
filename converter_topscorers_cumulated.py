"""
Amateurvoetbal: gecumuleerde topscorers (oude stand + nieuwe speelronde) -> Word (.docx)

Input:
- Bronbestand (Word of Kladblok): .txt / .docx / .doc
  Dit bestand bevat de actuele stand tot en met de vorige speelronde.
- Uitslagenbestand (Excel): .xlsx / .xls
  Dit bestand bevat de uitslagen van de huidige speelronde (tabblad 'INVOER').

Regels:
- Spelers matchen op naam (case-insensitive).
- Bij match: spelling/naam uit bronbestand blijft leidend in de output.
- Nieuwe spelers (niet in bron): spelling uit Excel wordt gebruikt.
- Outputopmaak is exact gelijk aan converter_topscorers.topscorers_text_to_docx_bytes
  (zelfde lijstnummering / koppen / vet nummer).

Let op:
- .doc (oud Word-formaat) wordt heuristisch uitgelezen. Als dat mislukt,
  geef een duidelijke foutmelding en gebruik bij voorkeur .docx.
"""

from __future__ import annotations

import io
import re
from collections import defaultdict
from dataclasses import dataclass
from typing import Dict, Tuple, Optional, List, Any

import openpyxl

try:
    import xlrd  # type: ignore
except Exception:  # pragma: no cover
    xlrd = None  # type: ignore

from docx import Document

from converter_topscorers import topscorers_text_to_docx_bytes


# ----------------------------
# Errors (technisch, geen inhoud)
# ----------------------------
@dataclass
class ConversionError(Exception):
    code: str
    message: str

    def __str__(self) -> str:
        return f"{self.code}: {self.message}"


# ----------------------------
# Limburgse clubs (uit notebook)
# ----------------------------
LIMBURG_CLUBS = {
    'Abdissenbosch',
    'Achates',
    'Alfa Sport',
    'America',
    'Amstenrade',
    'BEVO',
    'BMR',
    'BSV Limburgia',
    "BVV'27",
    'Baarlo',
    'Bekkerveld',
    'Belfeldia',
    "Berg'28",
    'Bieslo',
    'Blerick',
    'Boekel Sport',
    'Born',
    'Brevendia',
    'Bunde',
    'Caesar',
    'Chevremont',
    "Conventus'03",
    'DBSV',
    'DES Swalmen',
    'DESM',
    'DEV-Arcen',
    "DFO'20",
    "DVC'16",
    'DVO',
    'Daalhof',
    'De Dem',
    'De Leeuw',
    'De Ster',
    'EMS',
    'EVV',
    'Eijsden',
    'Eikenderveld',
    'Eindse Boys',
    'FC Bemelen',
    'FC Geleen Zuid',
    'FC Gulpen',
    'FC Hoensbroek',
    'FC Kerkrade-West',
    'FC Maasgouw',
    'FC ODA',
    'FC RIA',
    'FC Roerdalen',
    'FCV-Venlo',
    'FSG',
    "GSV'28",
    'Geertruidse Boys',
    'Geulsche Boys',
    'Geusselt Sport',
    'Grashoek',
    'Groene Ster',
    'H.B.S.V.',
    'Haelen',
    'Haslou',
    'Heer',
    'Hegelsom',
    'Heijen',
    'Helden',
    'Hellas',
    'Holthees-Smakt',
    'IVO',
    'IVS',
    'KSV Horn',
    'KVC Oranje',
    'Kakertse Boys',
    'Keer',
    'Koningslust',
    'Kronenberg',
    'Langeberg',
    'Leonidas-W',
    'Leunen',
    'Leveroy',
    'Lindenheuvel-Heidebloem Combinatie',
    'Linne',
    "Lottum-GFC'33",
    "MBC'13",
    'MMC Weert',
    'MSH Maasduinen',
    "MVC'19",
    'Melderslo',
    'Merefeldia',
    'Merselo',
    'Meterik',
    'Milsbeek',
    'Minor',
    'Neerbeek',
    'Oostrum',
    "PEC'20",
    'Partij',
    'Passart-VKC',
    "RIOS'31",
    'RKASV',
    'RKAVC',
    'RKDSO',
    'RKHBS',
    'RKHSV',
    'RKIVV',
    'RKMSV',
    'RKMVC',
    'RKSVB',
    'RKSVN',
    'RKSVO',
    'RKSVV',
    'RKTSV',
    'RKUVC',
    'RKVB',
    'RKVVM',
    'RVU',
    'Reuver',
    'Rimburg',
    'Roggel',
    "Rood Groen LVC'01",
    'Roosteren',
    'SCG',
    'SHH',
    'SNA',
    "SNC'14",
    "SSS'18",
    'SV Brunssum',
    'SV Geuldal',
    'SV Heythuysen',
    'SV Hulsberg',
    'SV Laar',
    'SV Meerssen',
    'SV Simpelveld',
    'SV United',
    'SV Venray',
    'SVC 2000',
    'SVEB-Sporting S.T.',
    "SVH'39",
    'SVM',
    'SVME',
    "SVOC'01",
    'Sanderbout',
    'Schaesberg',
    'Scharn',
    'Schimmert',
    'Schinveld',
    'Sittard',
    'Slekker Boys',
    "Sparta'18",
    'Spaubeek',
    'Sportclub Jekerdal',
    'Sportclub Leeuwen',
    'Sportclub Susteren',
    "Sportclub'25",
    'Sporting H.A.C.',
    'Sporting Heerlen',
    'St. Joost',
    'TSC Irene',
    "UOW'02",
    'Urmondia',
    'VV Hebes',
    'VV Kessel',
    'VV Maastricht West',
    'VV Schaesberg',
    'Vaesrade',
    'Venlosche Boys',
    'Veritas',
    'Vijlen',
    "Vitesse'08",
    'Voerendaal',
    'Walram',
    'Weltania',
    'Wijnandia',
    'Willem I',
    'Wittenhorst',
    'Woander Forest',
    'Ysselsteyn',
    "Zwart-Wit'19",
    'Zwentibold'
}


# ----------------------------
# Helpers: tekst uit uploads
# ----------------------------
def _decode_text_best_effort(raw: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return raw.decode(enc)
        except Exception:
            continue
    return raw.decode("utf-8", errors="replace")


def _strip_rtf(text: str) -> str:
    # Simpele RTF-stripper: voldoende voor "platte" RTF exports.
    # Verwijdert control words en groepen; houdt leesbare tekst over.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Verwijder rtf-headers
    text = re.sub(r"^\{\\rtf1.*?\n", "", text, flags=re.DOTALL)
    # Verwijder unicode escapes \uNNNN?
    text = re.sub(r"\\u-?\d+\??", "", text)
    # Verwijder control words \wordN
    text = re.sub(r"\\[a-zA-Z]+\d* ?", "", text)
    # Verwijder overige braces
    text = text.replace("{", "").replace("}", "")
    # Normaliseer whitespace
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_text_from_doc_heuristic(raw: bytes) -> str:
    # 1) RTF in .doc container komt voor
    if raw.startswith(b"{\\rtf"):
        return _strip_rtf(_decode_text_best_effort(raw))

    # 2) Heuristiek: probeer UTF-16LE en cp1252; kies kandidaat met meeste "signaal"
    candidates: List[str] = []
    try:
        candidates.append(raw.decode("utf-16le", errors="ignore"))
    except Exception:
        pass
    try:
        candidates.append(raw.decode("cp1252", errors="ignore"))
    except Exception:
        pass

    def score(s: str) -> int:
        sl = s.lower()
        return (
            sl.count("klasse") * 10
            + sl.count("divisie") * 10
            + sl.count("doelpunt") * 4
            + sl.count("\n")
        )

    best = ""
    best_score = -1
    for c in candidates:
        # haal extreem veel nulls/rommel weg
        c = c.replace("\x00", "")
        c = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\u02FF\u1E00-\u1EFF]", "", c)
        c = c.replace("\r\n", "\n").replace("\r", "\n")
        sc = score(c)
        if sc > best_score:
            best = c
            best_score = sc

    best = re.sub(r"\n{3,}", "\n\n", best).strip()
    if not best or best_score < 5:
        raise ConversionError(
            "TS-CUM-003",
            "Kon .doc niet betrouwbaar uitlezen. Sla het bestand op als .docx en upload opnieuw.",
        )
    return best


def extract_text_from_source_upload(raw: bytes, filename: str) -> str:
    name = (filename or "").lower()
    if name.endswith(".docx"):
        doc = Document(io.BytesIO(raw))
        lines: List[str] = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                lines.append(t)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for ln in cell.text.splitlines():
                        t = ln.strip()
                        if t:
                            lines.append(t)
        return "\n".join(lines)

    if name.endswith(".doc"):
        return _extract_text_from_doc_heuristic(raw)

    # .txt of onbekend: best-effort decode
    return _decode_text_best_effort(raw).strip()


# ----------------------------
# Parsing & samenvoegen (uit notebook, aangepast voor case-insensitive naam-match)
# ----------------------------
LINE_PATTERN = re.compile(
    r"^\s*(?:(\d+)\.\s*)?(.+?)(?:\s*-\s*(\d+)\s+doelpunt(?:en)?)?\s*$",
    re.IGNORECASE,
)


def canonical_group(header: str) -> str:
    h = header.strip()
    hl = h.lower()

    if "derde en vierde divisie" in hl:
        return "Derde en vierde divisie"
    if hl.startswith("eerste klasse"):
        return "Eerste klasse"
    if hl.startswith("tweede klasse"):
        return "Tweede klasse"
    if hl.startswith("derde klasse"):
        return "Derde klasse"
    if hl.startswith("vierde klasse"):
        return "Vierde klasse"
    if hl.startswith("vijfde klasse"):
        return "Vijfde klasse"
    return h


def split_name_club_round(raw_name_club: str) -> Tuple[str, Optional[str]]:
    raw_name_club = raw_name_club.strip()
    i = raw_name_club.rfind("(")
    j = raw_name_club.rfind(")")
    if i == -1 or j == -1 or j < i:
        return raw_name_club.strip(), None
    name = raw_name_club[:i].strip()
    club = raw_name_club[i + 1 : j].strip()
    return name, club


def split_name_club_bron(raw_name_club: str) -> Tuple[str, Optional[str], Optional[str]]:
    raw_name_club = raw_name_club.strip()
    i = raw_name_club.rfind("(")
    j = raw_name_club.rfind(")")
    if i == -1 or j == -1 or j < i:
        return raw_name_club.strip(), None, None
    name = raw_name_club[:i].strip()
    inside = raw_name_club[i + 1 : j].strip()
    parts = [p.strip() for p in inside.split(",")]
    club = parts[0] if parts else None
    extra = ", ".join(parts[1:]) or None
    return name, club, extra


def _norm_name(name: str) -> str:
    return re.sub(r"\s+", " ", (name or "").strip()).casefold()


def parse_totals_text(text: str) -> Tuple[Dict[Tuple[str, Optional[str]], int], Dict[Tuple[str, Optional[str]], Dict[str, Any]]]:
    totals_before: Dict[Tuple[str, Optional[str]], int] = {}
    meta_before: Dict[Tuple[str, Optional[str]], Dict[str, Any]] = {}
    current_group: Optional[str] = None
    last_goals_in_block: Optional[int] = None

    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue

        if "(" not in line:
            current_group = canonical_group(line)
            last_goals_in_block = None
            continue

        m = LINE_PATTERN.match(line)
        if not m:
            continue

        _, raw_name_club, goals_str = m.groups()
        name, club, extra = split_name_club_bron(raw_name_club)

        if goals_str is not None:
            last_goals_in_block = int(goals_str)

        if last_goals_in_block is None:
            continue

        key = (name, club)
        totals_before[key] = last_goals_in_block
        meta_before[key] = {"group": current_group, "extra": extra}

    return totals_before, meta_before


# ----------------------------
# Excel parsing (uit notebook)
# ----------------------------
SUSPICIOUS_NAMES = set()


def normalize_division_name(division):
    if not isinstance(division, str):
        return division
    up = division.upper()
    if "DERDE DIVISIE" in up or "VIERDE DIVISIE" in up:
        return "Derde en vierde divisie"
    return division


def parse_goals_cell(text, home_team, away_team):
    global SUSPICIOUS_NAMES
    if not text:
        return []

    pattern = re.compile(r"([^,\.]+?)\s+(\d+-\d+)")
    matches = list(pattern.finditer(text))

    results = []
    last_player_name = None
    last_end = 0
    prev_home, prev_away = 0, 0

    for match in matches:
        raw_name = match.group(1).strip()
        score = match.group(2)

        context = text[last_end : match.end(2)].lower()
        segment = text[match.start(1) : match.end(2)].lower()
        last_end = match.end(2)

        if (
            "eigen doelpunt" in context
            or "eigen doelpunt" in segment
            or "ed." in context
            or "ed." in segment
        ):
            try:
                home_goals, away_goals = map(int, score.split("-"))
                prev_home, prev_away = home_goals, away_goals
            except ValueError:
                SUSPICIOUS_NAMES.add(f"bad-score-eo:{score}")
            continue

        if isinstance(home_team, str) and isinstance(away_team, str):
            if home_team and away_team and home_team in raw_name and away_team in raw_name:
                continue

        name = raw_name

        while name.startswith("(") and ")" in name:
            name = name[name.find(")") + 1 :].strip()

        lower_name = name.lower()

        if lower_name == "en" or re.fullmatch(r"\d+-\d+\s+en", lower_name):
            if last_player_name is None:
                SUSPICIOUS_NAMES.add(f"en-without-previous:{name}")
                continue
            name = last_player_name
        else:
            if lower_name.startswith("en "):
                name = name[3:].strip()
            last_player_name = name

        if not any(c.isalpha() for c in name) or len(name) <= 1:
            SUSPICIOUS_NAMES.add(name)

        try:
            home_goals, away_goals = map(int, score.split("-"))
        except ValueError:
            SUSPICIOUS_NAMES.add(f"bad-score:{score}")
            continue

        # Negeer naam 'onbekend' (case-insensitive) in de telling, maar laat de stand doorlopen.
        if _norm_name(name) == "onbekend":
            prev_home, prev_away = home_goals, away_goals
            continue

        scored_side = None
        steps = 0

        if home_goals == prev_home + 1 and away_goals == prev_away:
            scored_side = "home"
            steps = 1
        elif away_goals == prev_away + 1 and home_goals == prev_home:
            scored_side = "away"
            steps = 1
        else:
            if name == last_player_name:
                if home_goals > prev_home and away_goals == prev_away:
                    scored_side = "home"
                    steps = home_goals - prev_home
                elif away_goals > prev_away and home_goals == prev_home:
                    scored_side = "away"
                    steps = away_goals - prev_away

        if scored_side is None or steps <= 0:
            SUSPICIOUS_NAMES.add(f"weird-score-seq:{prev_home}-{prev_away}->{home_goals}-{away_goals}")
            prev_home, prev_away = home_goals, away_goals
            continue

        club = home_team if scored_side == "home" else away_team
        for _ in range(steps):
            results.append((name, club))

        prev_home, prev_away = home_goals, away_goals

    return results


def _is_beker_marker(val: Any) -> bool:
    """True als deze rij het begin is van de beker-sectie (bijv. cel B274 = 'BEKER')."""
    if not isinstance(val, str):
        return False
    return val.strip().lower() == "beker"


def build_rankings(ws):
    max_row = ws.max_row
    rankings = {}

    header_rows = []
    for r in range(1, max_row + 1):
        # Stop met zoeken naar divisie-tabellen zodra de beker-sectie start.
        if _is_beker_marker(ws.cell(r, 2).value):
            break
        if (
            isinstance(ws.cell(r, 2).value, str)
            and ws.cell(r, 2).value
            and ws.cell(r, 6).value == "EINDSTAND"
        ):
            header_rows.append(r)

    for header in header_rows:
        raw_division = str(ws.cell(header, 2).value).strip()
        division = normalize_division_name(raw_division)
        players = rankings.setdefault(division, defaultdict(int))

        row = header + 1
        while row <= max_row:
            # Stop bij de beker-sectie (wordt niet meegenomen in de competitiestanden).
            if _is_beker_marker(ws.cell(row, 2).value):
                break
            home = ws.cell(row, 2).value
            away = ws.cell(row, 4).value

            if not home and not away:
                break

            scorers = ws.cell(row, 12).value
            if scorers:
                for name, club in parse_goals_cell(str(scorers), str(home), str(away)):
                    players[(name, club)] += 1
            row += 1

        rankings[division] = players

    return rankings


class _XlsSheetAdapter:
    """Adapter zodat de notebook-logica ook met xlrd (.xls) werkt (1-indexed cell access)."""

    def __init__(self, sheet):
        self._sheet = sheet
        self.max_row = int(getattr(sheet, "nrows", 0))

    class _Cell:
        def __init__(self, value):
            self.value = value

    def cell(self, r: int, c: int):
        rr = r - 1
        cc = c - 1
        try:
            val = self._sheet.cell_value(rr, cc)
        except Exception:
            val = None
        if val == "":
            val = None
        return self._Cell(val)


def parse_excel_round(excel_raw: bytes, excel_filename: str) -> Tuple[Dict[Tuple[str, Optional[str]], int], Dict[Tuple[str, Optional[str]], str]]:
    name = (excel_filename or "").lower()
    try:
        if name.endswith(".xls"):
            if xlrd is None:
                raise ConversionError("TS-CUM-004", "xlrd ontbreekt voor .xls. Upload een .xlsx-bestand.")
            book = xlrd.open_workbook(file_contents=excel_raw)
            try:
                sheet = book.sheet_by_name("INVOER")
            except Exception:
                raise ConversionError("TS-CUM-004", "Tabblad 'INVOER' ontbreekt in het Excelbestand.")
            ws = _XlsSheetAdapter(sheet)
        else:
            wb = openpyxl.load_workbook(io.BytesIO(excel_raw), data_only=True)
            if "INVOER" not in wb.sheetnames:
                raise ConversionError("TS-CUM-004", "Tabblad 'INVOER' ontbreekt in het Excelbestand.")
            ws = wb["INVOER"]
    except ConversionError:
        raise
    except Exception as e:
        raise ConversionError("TS-CUM-004", f"Kon Excelbestand niet openen: {e}")

    rankings = build_rankings(ws)

    goals_this_round: Dict[Tuple[str, Optional[str]], int] = {}
    groups_for_new: Dict[Tuple[str, Optional[str]], str] = {}

    for division, players_dict in rankings.items():
        group = canonical_group(str(division))
        for (pname, pclub), goals in players_dict.items():
            if not pclub or pclub not in LIMBURG_CLUBS:
                continue
            key = (str(pname).strip(), str(pclub).strip())
            goals_this_round[key] = goals_this_round.get(key, 0) + int(goals)
            if key not in groups_for_new:
                groups_for_new[key] = group

    return goals_this_round, groups_for_new


# ----------------------------
# Merge (naam case-insensitive; bron-spelling leidend)
# ----------------------------
def merge_totals_case_insensitive(
    totals_before: Dict[Tuple[str, Optional[str]], int],
    meta_before: Dict[Tuple[str, Optional[str]], Dict[str, Any]],
    goals_this_round: Dict[Tuple[str, Optional[str]], int],
    groups_for_new: Dict[Tuple[str, Optional[str]], str],
):
    """Voeg 'oude stand' + 'huidige ronde' samen op (naam, club).

    - Naam-matching is case-insensitive (casefold) binnen dezelfde club.
    - Club-matching is case-insensitive (casefold).
    - Spelling uit de bron blijft leidend wanneer er een match is.
    - Naam 'onbekend' wordt genegeerd in de telling.
    """

    # Filter 'onbekend' uit bron (voor het geval het erin staat)
    totals_before = {
        k: v for k, v in totals_before.items() if _norm_name(k[0]) != "onbekend"
    }
    meta_before = {
        k: v for k, v in meta_before.items() if _norm_name(k[0]) != "onbekend"
    }

    def _norm_club(club: Optional[str]) -> str:
        return _norm_name(club or "")

    # Index bron op (genormaliseerde naam, genormaliseerde club)
    idx_before: Dict[Tuple[str, str], Tuple[str, Optional[str]]] = {}
    for key in totals_before.keys():
        nk = (_norm_name(key[0]), _norm_club(key[1]))
        if nk in idx_before:
            # Zelfde speler (na normalisatie) komt dubbel voor binnen dezelfde club
            raise ConversionError(
                "TS-CUM-005",
                "De bron-stand bevat dezelfde speler meerdere keren bij dezelfde club. Pas dit aan en probeer opnieuw.",
            )
        idx_before[nk] = key

    totals_after: Dict[Tuple[str, Optional[str]], Dict[str, Any]] = {}

    # Start met bron-stand
    for key, total in totals_before.items():
        meta = meta_before.get(key, {})
        totals_after[key] = {
            "goals": int(total),
            "group": meta.get("group"),
            "extra": meta.get("extra"),
        }

    # Groepen-index (voor fallback bij case-insensitive clubnaam)
    idx_groups: Dict[Tuple[str, str], str] = {}
    for (n, c), g in groups_for_new.items():
        idx_groups[(_norm_name(n), _norm_club(c))] = g

    # Voeg ronde toe
    for (name_round, club_round), extra_goals in goals_this_round.items():
        if _norm_name(name_round) == "onbekend":
            continue

        nk = (_norm_name(name_round), _norm_club(club_round))

        target_key = idx_before.get(nk)
        if target_key:
            totals_after[target_key]["goals"] += int(extra_goals)
            continue

        group = groups_for_new.get((name_round, club_round)) or idx_groups.get(nk) or "Overig"
        totals_after[(name_round, club_round)] = {
            "goals": int(extra_goals),
            "group": group,
            "extra": None,
        }

    return totals_after



# ----------------------------
# Output: bron-stand tekst + docx
# ----------------------------
def build_class_block(klassenaam: str, players: List[Dict[str, Any]]) -> List[str]:
    players_sorted = sorted(
        players,
        key=lambda p: (-p["goals"], p["club"] or "", p["name"]),
    )

    lines = [klassenaam, ""]
    last_goals = None
    rank_counter = 0

    for p in players_sorted:
        goals = int(p["goals"])
        name = p["name"]
        club = p.get("club")
        extra = p.get("extra")

        if club and extra:
            inside = f"{club}, {extra}"
        elif club:
            inside = f"{club}"
        else:
            inside = ""

        parens = f" ({inside})" if inside else ""

        if goals != last_goals:
            rank_counter += 1
            last_goals = goals
            doelpunt_woord = "doelpunt" if goals == 1 else "doelpunten"
            line = f"{rank_counter}. {name}{parens} - {goals} {doelpunt_woord}"
        else:
            line = f"{name}{parens}"

        lines.append(line)

    return lines


def build_new_totals_text(totals_after: Dict[Tuple[str, Optional[str]], Dict[str, Any]]) -> str:
    per_class: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
    for (name, club), info in totals_after.items():
        group = info.get("group") or "Overig"
        per_class[str(group)].append(
            {
                "name": name,
                "club": club,
                "goals": int(info["goals"]),
                "extra": info.get("extra"),
            }
        )

    preferred_order = [
        "Derde en vierde divisie",
        "Eerste klasse",
        "Tweede klasse",
        "Derde klasse",
        "Vierde klasse",
        "Vijfde klasse",
    ]

    ordered_classes: List[str] = []
    seen = set()

    for g in preferred_order:
        if g in per_class:
            ordered_classes.append(g)
            seen.add(g)

    for g in sorted(per_class.keys()):
        if g not in seen:
            ordered_classes.append(g)
            seen.add(g)

    all_lines: List[str] = []
    for klassenaam in ordered_classes:
        all_lines.extend(build_class_block(klassenaam, per_class[klassenaam]))
        all_lines.append("")

    return "\n".join(all_lines).strip() + "\n"


def cumulated_topscorers_to_docx_bytes(
    source_raw: bytes,
    source_filename: str,
    excel_raw: bytes,
    excel_filename: str,
) -> bytes:
    try:
        source_text = extract_text_from_source_upload(source_raw, source_filename)
        totals_before, meta_before = parse_totals_text(source_text)

        goals_this_round, groups_for_new = parse_excel_round(excel_raw, excel_filename)

        totals_after = merge_totals_case_insensitive(
            totals_before, meta_before, goals_this_round, groups_for_new
        )

        new_text = build_new_totals_text(totals_after)
        return topscorers_text_to_docx_bytes(new_text)
    except ConversionError:
        raise
    except Exception as e:
        raise ConversionError("TS-CUM-007", f"Onverwachte fout tijdens verwerken: {e}")